import openai
from docx import Document
from moviepy.editor import VideoFileClip

def trancribe_audio(audio_path):
    with open(audio_path, 'rb') as audio_file:
        transcription = openai.Audio.transcribe("whisper-1", audio_file)
    return transcription['text']

def saves_to_docx(minutes, filename):
    doc = Document()
    for k, v in minutes.items():
        heading = ' '.join(word.capitalize() for word in k.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(v)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)

def extract_audio_from_video(video_path):
    # Load the video file
    clip = VideoFileClip(video_path)

    # Extract the audio and save it as an MP3 file
    audio = clip.audio
    audio.write_audiofile("audio.mp3")

    # Close the video file
    clip.close()
    return "audio.mp3"
    
def main():
    # set up the api key
    openai.api_key_path = "api_key.key"
    # ask [1] for video or [2] for audio
    choice = input("Enter [1] for video or [2] for audio: ")
    if choice == '1':
        # ask the path to the user
        video_path = input("Enter the path to the video file: ")
        audio_path = extract_audio_from_video(video_path)
    elif choice == '2':
        # ask the path to the user
        audio_path = input("Enter the path to the audio file: ") 
    
    transcripton = trancribe_audio(audio_path)
    # ask for the filename
    filename = input("Enter the filename: ")
    saves_to_docx(transcripton, filename+'.docx')
    
if __name__ == '__main__':
    main()